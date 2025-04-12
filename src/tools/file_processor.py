import os
import pandas as pd
from typing import Dict, Any, List, Optional
import mimetypes
from pathlib import Path
import json
import logging

# For PDF processing
from pypdf import PdfReader
# For DOCX processing
import docx

from src.tools.tool_registry import Tool, tool_registry
from src.config.config import UPLOAD_FOLDER, ALLOWED_EXTENSIONS

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileProcessor(Tool):
    """Tool for processing different types of files."""
    
    name = "FileProcessor"
    description = "Process and extract data from different file types (CSV, Excel, PDF, DOCX)"
    
    input_schema = {
        "type": "object",
        "properties": {
            "file_path": {
                "type": "string",
                "description": "Path to the file to process"
            },
            "file_type": {
                "type": "string",
                "description": "Type of file (csv, excel, pdf, docx). If not provided, it will be inferred."
            },
            "sheet_name": {
                "type": "string",
                "description": "For Excel files, the name of the sheet to process"
            },
            "extract_tables": {
                "type": "boolean",
                "description": "For PDF and DOCX, whether to extract tables"
            }
        },
        "required": ["file_path"]
    }
    
    output_schema = {
        "type": "object",
        "properties": {
            "content_type": {
                "type": "string",
                "description": "Type of the extracted content"
            },
            "text_content": {
                "type": "string",
                "description": "Extracted text content if applicable"
            },
            "table_data": {
                "type": "array",
                "description": "Extracted tabular data if applicable"
            },
            "metadata": {
                "type": "object",
                "description": "File metadata"
            }
        }
    }
    
    def __init__(self):
        super().__init__()
    
    def execute(self, file_path: str, file_type: Optional[str] = None, 
                sheet_name: Optional[str] = None, extract_tables: bool = True) -> Dict[str, Any]:
        """
        Process a file and extract its contents.
        
        Args:
            file_path: Path to the file to process
            file_type: Type of file, if not provided will be inferred
            sheet_name: For Excel files, the name of the sheet to process
            extract_tables: For PDF and DOCX, whether to extract tables
            
        Returns:
            Dictionary with processed content
        """
        try:
            # Ensure file exists
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Get file type if not provided
            if not file_type:
                file_type = self._infer_file_type(file_path)
            
            # Handle binary files safely
            try:
                is_binary = self._is_binary_file(file_path)
                if is_binary and file_type not in ['pdf', 'docx', 'doc', 'xlsx', 'xls']:
                    return {
                        "content": f"Binary file detected: {file_path}",
                        "content_type": "binary",
                        "tables": [],
                        "metadata": {
                            "type": file_type,
                            "file_name": os.path.basename(file_path),
                            "is_binary": True,
                            "file_size": os.path.getsize(file_path)
                        }
                    }
            except Exception as e:
                logger.warning(f"Could not check if file is binary: {str(e)}")
            
            # Process based on file type
            if file_type == 'csv':
                return self._process_csv(file_path)
            elif file_type in ['xlsx', 'xls']:
                return self._process_excel(file_path, sheet_name)
            elif file_type == 'pdf':
                return self._process_pdf(file_path, extract_tables)
            elif file_type in ['docx', 'doc']:
                return self._process_docx(file_path, extract_tables)
            elif file_type == 'txt':
                return self._process_text(file_path)
            else:
                # Try to process as text if unsure
                try:
                    return self._process_text(file_path)
                except:
                    return {
                        "content": f"Unsupported file type: {file_type}",
                        "content_type": "error",
                        "tables": [],
                        "metadata": {"type": file_type, "error": "Unsupported file type"}
                    }
                
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            return {
                "content": f"Error processing file: {str(e)}",
                "content_type": "error",
                "tables": [],
                "metadata": {"type": "unknown", "error": str(e)}
            }
    
    def _infer_file_type(self, file_path: str) -> str:
        """Infer the file type from the file extension."""
        _, ext = os.path.splitext(file_path)
        if ext:
            return ext[1:].lower()  # Remove the dot
        else:
            # Try to guess from content
            mime_type, _ = mimetypes.guess_type(file_path)
            if mime_type:
                for ext, mime in ALLOWED_EXTENSIONS.items():
                    if mime == mime_type:
                        return ext
            
            raise ValueError(f"Could not determine file type for: {file_path}")
    
    def _process_csv(self, file_path: str) -> Dict[str, Any]:
        """Process a CSV file."""
        try:
            df = pd.read_csv(file_path)
            
            # Basic statistics
            stats = {
                "row_count": len(df),
                "column_count": len(df.columns),
                "columns": df.columns.tolist()
            }
            
            # Convert to string for content preview
            content_preview = df.head(10).to_string(index=False)
            
            return {
                "content": content_preview,
                "content_type": "tabular",
                "tables": df.head(100).to_dict(orient='records'),  # Limit to first 100 rows
                "metadata": {
                    "type": "csv",
                    "file_name": os.path.basename(file_path),
                    "stats": stats,
                    "file_size": os.path.getsize(file_path)
                }
            }
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {str(e)}")
            return {
                "content": f"Error processing CSV: {str(e)}",
                "content_type": "error",
                "tables": [],
                "metadata": {"type": "csv", "error": str(e)}
            }
    
    def _process_excel(self, file_path: str, sheet_name: Optional[str] = None) -> Dict[str, Any]:
        """Process an Excel file."""
        try:
            if sheet_name:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheet_names = [sheet_name]
            else:
                # Read all sheets
                xlsx = pd.ExcelFile(file_path)
                sheet_names = xlsx.sheet_names
                
                # Read the first sheet if no specific sheet requested
                df = pd.read_excel(file_path, sheet_name=sheet_names[0])
            
            # Basic statistics
            stats = {
                "row_count": len(df),
                "column_count": len(df.columns),
                "columns": df.columns.tolist(),
                "sheets": sheet_names
            }
            
            # Convert to string for content preview
            content_preview = df.head(10).to_string(index=False)
            
            return {
                "content": content_preview,
                "content_type": "tabular",
                "tables": df.head(100).to_dict(orient='records'),  # Limit to first 100 rows
                "metadata": {
                    "type": "excel",
                    "file_name": os.path.basename(file_path),
                    "sheet_name": sheet_name if sheet_name else sheet_names[0],
                    "available_sheets": sheet_names,
                    "stats": stats,
                    "file_size": os.path.getsize(file_path)
                }
            }
        except Exception as e:
            logger.error(f"Error processing Excel {file_path}: {str(e)}")
            return {
                "content": f"Error processing Excel: {str(e)}",
                "content_type": "error",
                "tables": [],
                "metadata": {"type": "excel", "error": str(e)}
            }
    
    def _process_pdf(self, file_path: str, extract_tables: bool = True) -> Dict[str, Any]:
        """Process a PDF file."""
        try:
            reader = PdfReader(file_path)
            num_pages = len(reader.pages)
            
            # Extract text
            text_content = ""
            for i in range(num_pages):
                page = reader.pages[i]
                page_text = page.extract_text()
                if page_text:
                    text_content += page_text + "\n\n"
                else:
                    text_content += f"[Page {i+1} contains no extractable text or only images]\n\n"
            
            # TODO: Implement table extraction from PDF if extract_tables is True
            # This would require additional libraries like tabula-py
            
            return {
                "content": text_content,
                "content_type": "document",
                "tables": [],  # Placeholder for table extraction
                "metadata": {
                    "type": "pdf",
                    "file_name": os.path.basename(file_path),
                    "num_pages": num_pages,
                    "file_size": os.path.getsize(file_path)
                }
            }
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            return {
                "content": f"Error processing PDF: {str(e)}",
                "content_type": "error",
                "tables": [],
                "metadata": {"type": "pdf", "error": str(e)}
            }
    
    def _process_docx(self, file_path: str, extract_tables: bool = True) -> Dict[str, Any]:
        """Process a DOCX file."""
        try:
            doc = docx.Document(file_path)
            
            # Extract text paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty paragraphs
                    paragraphs.append(para.text)
            
            text_content = "\n".join(paragraphs)
            
            # Extract tables if requested
            tables = []
            if extract_tables and doc.tables:
                for i, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    tables.append({
                        "table_index": i,
                        "data": table_data
                    })
            
            return {
                "content": text_content,
                "content_type": "document",
                "tables": tables,
                "metadata": {
                    "type": "docx",
                    "file_name": os.path.basename(file_path),
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                    "file_size": os.path.getsize(file_path)
                }
            }
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            return {
                "content": f"Error processing DOCX: {str(e)}",
                "content_type": "error",
                "tables": [],
                "metadata": {"type": "docx", "error": str(e)}
            }
    
    def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Process a plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            return {
                "content": text_content,
                "content_type": "text",
                "tables": [],
                "metadata": {
                    "type": "txt",
                    "file_name": os.path.basename(file_path),
                    "file_size": os.path.getsize(file_path),
                    "line_count": text_content.count('\n') + 1
                }
            }
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {str(e)}")
            return {
                "content": f"Error processing text file: {str(e)}",
                "content_type": "error",
                "tables": [],
                "metadata": {"type": "txt", "error": str(e)}
            }

    def _is_binary_file(self, file_path: str, check_size: int = 1024) -> bool:
        """
        Check if a file is binary by reading the first chunk
        and looking for null bytes and unprintable characters.
        """
        with open(file_path, 'rb') as f:
            chunk = f.read(check_size)
            
        # Check for null bytes (common in binary files)
        if b'\x00' in chunk:
            return True
            
        # Try to decode as text
        try:
            chunk.decode('utf-8')
            return False
        except UnicodeDecodeError:
            # If it can't be decoded as UTF-8, it's likely binary
            return True

# Register the tool
tool_registry.register(FileProcessor) 