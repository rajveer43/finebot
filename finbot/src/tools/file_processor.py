import os
import pandas as pd
from typing import Dict, Any, List, Optional
import mimetypes
from pathlib import Path
import json
import logging

# For PDF processing
from pypdf import PdfReader
# For DOCX processing
import docx

from src.tools.tool_registry import Tool, tool_registry
from src.config.config import UPLOAD_FOLDER, ALLOWED_EXTENSIONS

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileProcessor(Tool):
    """Tool for processing different types of files."""
    
    name = "FileProcessor"
    description = "Process and extract data from different file types (CSV, Excel, PDF, DOCX)"
    
    input_schema = {
        "type": "object",
        "properties": {
            "file_path": {
                "type": "string",
                "description": "Path to the file to process"
            },
            "file_type": {
                "type": "string",
                "description": "Type of file (csv, excel, pdf, docx). If not provided, it will be inferred."
            },
            "sheet_name": {
                "type": "string",
                "description": "For Excel files, the name of the sheet to process"
            },
            "extract_tables": {
                "type": "boolean",
                "description": "For PDF and DOCX, whether to extract tables"
            }
        },
        "required": ["file_path"]
    }
    
    output_schema = {
        "type": "object",
        "properties": {
            "content_type": {
                "type": "string",
                "description": "Type of the extracted content"
            },
            "text_content": {
                "type": "string",
                "description": "Extracted text content if applicable"
            },
            "table_data": {
                "type": "array",
                "description": "Extracted tabular data if applicable"
            },
            "metadata": {
                "type": "object",
                "description": "File metadata"
            }
        }
    }
    
    def __init__(self):
        super().__init__()
    
    def execute(self, file_path: str, file_type: Optional[str] = None, 
                sheet_name: Optional[str] = None, extract_tables: bool = True) -> Dict[str, Any]:
        """
        Process a file and extract its contents.
        
        Args:
            file_path: Path to the file to process
            file_type: Type of file, if not provided will be inferred
            sheet_name: For Excel files, the name of the sheet to process
            extract_tables: For PDF and DOCX, whether to extract tables
            
        Returns:
            Dictionary with processed content
        """
        try:
            # Ensure file exists
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Get file type if not provided
            if not file_type:
                file_type = self._infer_file_type(file_path)
            
            # Process based on file type
            if file_type == 'csv':
                return self._process_csv(file_path)
            elif file_type in ['xlsx', 'xls']:
                return self._process_excel(file_path, sheet_name)
            elif file_type == 'pdf':
                return self._process_pdf(file_path, extract_tables)
            elif file_type in ['docx', 'doc']:
                return self._process_docx(file_path, extract_tables)
            elif file_type == 'txt':
                return self._process_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            return {
                "error": str(e),
                "content_type": "error",
                "file_path": file_path
            }
    
    def _infer_file_type(self, file_path: str) -> str:
        """Infer the file type from the file extension."""
        _, ext = os.path.splitext(file_path)
        if ext:
            return ext[1:].lower()  # Remove the dot
        else:
            # Try to guess from content
            mime_type, _ = mimetypes.guess_type(file_path)
            if mime_type:
                for ext, mime in ALLOWED_EXTENSIONS.items():
                    if mime == mime_type:
                        return ext
            
            raise ValueError(f"Could not determine file type for: {file_path}")
    
    def _process_csv(self, file_path: str) -> Dict[str, Any]:
        """Process a CSV file."""
        df = pd.read_csv(file_path)
        
        # Basic statistics
        stats = {
            "row_count": len(df),
            "column_count": len(df.columns),
            "columns": df.columns.tolist()
        }
        
        return {
            "content_type": "tabular",
            "text_content": "",
            "table_data": df.head(100).to_dict(orient='records'),  # Limit to first 100 rows
            "metadata": {
                "file_type": "csv",
                "file_name": os.path.basename(file_path),
                "stats": stats
            }
        }
    
    def _process_excel(self, file_path: str, sheet_name: Optional[str] = None) -> Dict[str, Any]:
        """Process an Excel file."""
        if sheet_name:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            sheet_names = [sheet_name]
        else:
            # Read all sheets
            xlsx = pd.ExcelFile(file_path)
            sheet_names = xlsx.sheet_names
            
            # Read the first sheet if no specific sheet requested
            df = pd.read_excel(file_path, sheet_name=sheet_names[0])
        
        # Basic statistics
        stats = {
            "row_count": len(df),
            "column_count": len(df.columns),
            "columns": df.columns.tolist(),
            "sheets": sheet_names
        }
        
        return {
            "content_type": "tabular",
            "text_content": "",
            "table_data": df.head(100).to_dict(orient='records'),  # Limit to first 100 rows
            "metadata": {
                "file_type": "excel",
                "file_name": os.path.basename(file_path),
                "sheet_name": sheet_name if sheet_name else sheet_names[0],
                "available_sheets": sheet_names,
                "stats": stats
            }
        }
    
    def _process_pdf(self, file_path: str, extract_tables: bool = True) -> Dict[str, Any]:
        """Process a PDF file."""
        reader = PdfReader(file_path)
        num_pages = len(reader.pages)
        
        # Extract text
        text_content = ""
        for i in range(num_pages):
            page = reader.pages[i]
            text_content += page.extract_text() + "\n\n"
        
        # TODO: Implement table extraction from PDF if extract_tables is True
        # This would require additional libraries like tabula-py
        
        return {
            "content_type": "document",
            "text_content": text_content,
            "table_data": [],  # Placeholder for table extraction
            "metadata": {
                "file_type": "pdf",
                "file_name": os.path.basename(file_path),
                "num_pages": num_pages
            }
        }
    
    def _process_docx(self, file_path: str, extract_tables: bool = True) -> Dict[str, Any]:
        """Process a DOCX file."""
        doc = docx.Document(file_path)
        
        # Extract text paragraphs
        text_content = "\n".join([para.text for para in doc.paragraphs])
        
        # Extract tables if requested
        tables = []
        if extract_tables and doc.tables:
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    "table_index": i,
                    "data": table_data
                })
        
        return {
            "content_type": "document",
            "text_content": text_content,
            "table_data": tables,
            "metadata": {
                "file_type": "docx",
                "file_name": os.path.basename(file_path),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
        }
    
    def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Process a plain text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            text_content = f.read()
        
        return {
            "content_type": "text",
            "text_content": text_content,
            "table_data": [],
            "metadata": {
                "file_type": "txt",
                "file_name": os.path.basename(file_path),
                "file_size": os.path.getsize(file_path)
            }
        }

# Register the tool
tool_registry.register(FileProcessor) 